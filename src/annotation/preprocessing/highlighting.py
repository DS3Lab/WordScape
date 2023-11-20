from docx.document import Document as DocxDocument
from docx.table import _Cell


def sanitize_highlighting(word_doc: DocxDocument) -> DocxDocument:
    r"""Remove highlighting from a word document, as this interferes with our
    colorization based annotation process.

    @param word_doc: word document instance

    @return: sanitized word document instance
    """
    for para in word_doc.paragraphs:
        # Iterate over all runs in the paragraph
        for run in para.runs:
            # Check if the run has highlighting
            if run.font.highlight_color is not None:
                # Remove the highlighting
                run.font.highlight_color = None

    # Iterate over all tables in the document
    for table in word_doc.tables:
        for row in table.rows:
            try:
                row_cells = row.cells
            except IndexError:
                row_cells = [_Cell(tc, table) for tc in row._tr.tc_lst]
            for cell in row_cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        # remove highlighting
                        if run.font.highlight_color is not None:
                            run.font.highlight_color = None

    return word_doc
