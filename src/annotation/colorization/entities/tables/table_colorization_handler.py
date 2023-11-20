from lxml import etree
from typing import Tuple, Callable, Union

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.styles import CT_Style
from docx.oxml.table import CT_Tbl, CT_Tc, CT_Row
from docx.shared import RGBColor
from docx.table import _Cell

from src.annotation.colorization import ColorizationHandler
from src.annotation.colorization.entities.tables import styles
from src.annotation.colorization.entities.tables.element_parsers import *
from src.annotation.colorization.entities.tables.utils import *
from src.annotation.utils.color_utils import rgb_to_hex, hsv_to_rgb

import settings


def _make_color_cycle_step(sat, sat_base, val, val_base, sat_val_step):
    r"""Make a step through the HSV color space. The function returns the new
    saturation and value.

    @param sat: The current saturation.
    @param sat_base: The base saturation value.
    @param val: The current val value.
    @param val_base: The base val value.
    @param sat_val_step: The step size used to make a step through the HSV
        color space (applies to both saturation and value).

    @return: The new saturation and value levels.
    """
    # decrease val
    val -= sat_val_step

    # check if val is below the minimum value; if so, reset value to the base
    # value level and decrease the saturation level.
    if val < settings.colors.VAL_MIN:
        # reset val
        val = val_base

        # decrease sat
        sat -= sat_val_step

        # check if sat is below the minimum value; if so, reset saturation to
        # the base saturation level.
        if sat < settings.colors.SAT_MIN:
            sat = sat_base  # reset sat

    return sat, val


class TableColorizationHandler:
    def __init__(
            self,
            ct_tbl: CT_Tbl,
            ct_tbl_ref_style: CT_Style,
            colorization_handler: ColorizationHandler,
            base_color_table: Tuple[int, int, int],
            base_color_header: Union[Tuple[int, int, int], None],
            sat_val_step: int = settings.colors.SAT_VAL_STEP
    ):
        self._ct_tbl = ct_tbl
        self._ct_tbl_ref_style = ct_tbl_ref_style
        self._colorization_handler = colorization_handler
        self._base_color_table = base_color_table
        self._base_color_header = base_color_header
        self._sat_val_step = sat_val_step

        # parse the table style referenced by the table. This defines 1)
        # conditional styles used in the table, and 2) values which are applied
        # where higher priority styles di not define them.
        self.__parse_tbl_ref_style()

        # parse the the style defined for the table itself. This defines both
        # conditional and non-conditional styles which override the styles
        # defined in the referenced style.
        self.__parse_table_properties()

    def colorize_table(self):
        # colorize header rows
        normal_row_idx, hsv_color_header = self.__colorize_rows(
            base_color=self._base_color_header,
            sat_val_step=self._sat_val_step,
            row_start_idx=0,
            breaking_condition=lambda row: not check_if_header_row(row)
        )

        # colorize normal rows
        _, hsv_color_table = self.__colorize_rows(
            base_color=self._base_color_table,
            sat_val_step=self._sat_val_step,
            row_start_idx=normal_row_idx,
            breaking_condition=lambda row: False
        )

        return hsv_color_header, hsv_color_table

    def __colorize_rows(
            self,
            base_color: Tuple[int, int, int],
            sat_val_step: int,
            row_start_idx: int,
            breaking_condition: Callable[[CT_Row], bool],
    ) -> Tuple[int, Union[Tuple[int, int, int], None]]:
        r""" Colorize rows of the table.

        Note: This method is used to colorize the header rows and the normal
            rows of the table. The function colorizes row until the breaking
            condition is met. The breaking condition is defined by the
            `breaking_condition` parameter. The function returns the index of
            the next row which has not yet been colorized.

        @param base_color: The base color used to colorize the rows in HSV
            color space.
        @param sat_val_step: The step size used to make a step through the
            HSV color space.
        @param row_start_idx: The index of the row where the colorization
            should start.
        @param breaking_condition: The function which defines the breaking
            condition. The function takes a row as input and returns a boolean
            value. If the function returns `True`, the colorization is stopped
            and the function returns the index of the current row.

        @return: The index of the next row which has not yet been colorized,
            and the HSV color of the last row which has been colorized.
        """
        if base_color is None:
            return 0, None

        hue, sat, val = base_color
        row_idx = row_start_idx

        for row in self._ct_tbl.tr_lst[row_start_idx:]:
            if breaking_condition(row):
                break

            for cell_num, tc in enumerate(row.tc_lst):
                # get cell location
                cell_loc = get_cell_location(
                    cell_num=cell_num, row_num=row_idx,
                    total_cells_in_row=len(row.tc_lst),
                    total_rows_in_table=len(self._ct_tbl.tr_lst)
                )

                cell_color_rgb = hsv_to_rgb(hsv_color=(hue, sat, val))
                cell_color_hex = rgb_to_hex(rgb_color=cell_color_rgb)

                self.__colorize_cell(
                    tc, row, hex_color=cell_color_hex,
                    rgb_color=cell_color_rgb, cell_loc=cell_loc
                )

                # record the color used to colorize the cell
                self._colorization_handler.update_application_color(
                    base_color=base_color, new_color=(hue, sat, val)
                )

                # make step through hsv color space
                sat, val = _make_color_cycle_step(
                    sat, base_color[1], val, base_color[2], sat_val_step
                )

                # colorize tables in cell
                for table in _Cell(tc, self._ct_tbl).tables:
                    try:
                        table_ref_style = table.style._element
                    except AttributeError:
                        table_ref_style = self._ct_tbl_ref_style

                    handler = TableColorizationHandler(
                        ct_tbl=table._tbl, ct_tbl_ref_style=table_ref_style,
                        colorization_handler=self._colorization_handler,
                        base_color_table=(hue, sat, val),
                        base_color_header=None,
                        sat_val_step=self._sat_val_step
                    )
                    _, (hue, sat, val) = handler.colorize_table()

            row_idx += 1

        return row_idx, (hue, sat, val)

    def __colorize_cell(
            self, tc: CT_Tc, tr: CT_Row, hex_color: str,
            rgb_color: Tuple[int, int, int], cell_loc: str
    ):
        r"""Colorize a cell in a table taking into account to the table styles
        defined in word/document.xml and word/styles.xml.

        @param tc: The cell to be colored
        @param tr: The row of the cell to be colored
        @param hex_color: The hex color to be applied to the cell
        @param rgb_color: The rgb color to be applied to the cell
        """
        tc_props_applied = self.__get_tc_props_applied(tc, tr, cell_loc)
        tc_borders = tc_props_applied.tc_borders
        tc_shd = tc_props_applied.tc_shd

        # apply shading to paragraphs in cell
        for par in _Cell(tc, self._ct_tbl).paragraphs:
            shade_element(
                par._p.get_or_add_pPr(), color_hex=hex_color, val=tc_shd.val
            )
            for run in par.runs:
                # adjust run level font color
                run.font.color.rgb = RGBColor(*rgb_color)

                # adjust run level shading if existing
                shade_element(
                    run._r.get_or_add_rPr(), color_hex=hex_color,
                    val=tc_shd.val
                )

        # add cell shading
        shade_element(
            tc.get_or_add_tcPr(), color_hex=hex_color, val=tc_shd.val
        )

        # Adjust cell border colors. This gets or adds a new tcBorders element
        # to the tcPr element of the cell. If the tcBorders element already
        # exists, the color of the existing borders is overwritten, and all
        # other attributes are kept taking into account the table style
        # hierarchy.
        tc_props = tc.get_or_add_tcPr()
        if (tc_borders_elem := tc_props.find(qn('w:tcBorders'))) is None:
            tc_borders_elem = OxmlElement('w:tcBorders')
            tc_props.append(tc_borders_elem)

        # iterate over all borders of the cell (top, left, bottom, right,
        # insideH, insideV, tl2br, tr2bl)
        for b in tc_borders.__slots__:
            if (elem_tag := tc_borders_elem.find(qn(f'w:{b}'))) is None:
                elem_tag = OxmlElement(f'w:{b}')
                tc_borders_elem.append(elem_tag)

            # iterate over all attributes of the current border (val, sz,
            # space, color) and overwrite the color
            for attr in getattr(tc_borders, b).__slots__:
                if (
                        attr_val := getattr(getattr(tc_borders, b), attr)
                ) is None:
                    continue

                attr_val = hex_color if attr == 'color' else attr_val
                elem_tag.set(qn(f'w:{attr}'), attr_val)

    def __get_tc_props_applied(
            self, tc: CT_Tc, tr: CT_Row, cell_loc: str
    ) -> styles.TableCellProperty:
        r"""Get the actual table cell properties applied to a cell taking into
        account the table style hierarchy.

        @param tc: The cell to be colored
        @param tr: The row of the cell to be colored

        @return: The table cell properties applied to the cell, consisting of
            shading and border styles
        """
        # create border style based on hierarchy
        tc_borders = self.__get_applied_tc_borders(tc, cell_loc=cell_loc)

        # create shading style based on hierarchy
        tc_shd = self.__get_applied_tc_shd(tc)

        # parse conditional formatting
        cond_style_applied = self.__get_tc_cond_style_applied(tc, tr)

        # consolidate conditional and non-conditional styles
        tc_borders = fill_border_style(
            borders=tc_borders, new_borders=cond_style_applied.tc_borders
        )
        tc_shd = fill_shd_style(
            shd=tc_shd, new_shd=cond_style_applied.tc_shd
        )

        return styles.TableCellProperty(
            tc_borders=tc_borders, tc_shd=tc_shd
        )

    def __get_tc_cond_style_applied(
            self, tc: CT_Tc, tr: CT_Row
    ) -> styles.TableCellProperty:
        r"""Get the conditional table cell properties applied to a cell taking
        into account the table style hierarchy.

        @param tc: The cell to be colored
        @param tr: The row of the cell to be colored

        @return: The conditional table cell properties applied to the cell,
            consisting of shading and border styles
        """
        cnf_style = styles.CnfStyle.init_zero()

        if not self._tbl_look.use_conditional_styling:
            # conditional styling is not applied
            return styles.TableCellProperty.init_zero()

        # cell level cnf style
        tc_etree = etree.fromstring(tc.xml)
        if (tc_pr := tc_etree.find(qn('w:tcPr'))) is not None:
            if (tc_cnf := tc_pr.find(qn('w:cnfStyle'))) is not None:
                for _attr in cnf_style.__slots__:
                    setattr(cnf_style, _attr, tc_cnf.get(qn(f'w:{_attr}')))

        # row level cnf style
        tr_etree = etree.fromstring(tr.xml)
        if (tr_pr := tr_etree.find(qn('w:trPr'))) is not None:
            if (tr_cnf := tr_pr.find(qn('w:cnfStyle'))) is not None:
                for _attr in cnf_style.__slots__:
                    if getattr(cnf_style, _attr) is not None:
                        continue
                    setattr(cnf_style, _attr, tr_cnf.get(qn(f'w:{_attr}')))

        # determine conditional style
        applied_styles = []

        # get conditional style for first row, if the current row is the first
        # row of the table and the table has a first row style defined
        if (
                cnf_style.firstRow in ["1", "true"] and
                self._tbl_look.firstRow in ["1", "true"]
        ):
            applied_styles.append(self._ref_cnd_styles.firstRow)

        # get conditional style for intermediate rows, if the current row is
        # neither the first nor the last row of the table and the table has
        # an intermediate row style defined. This applies to even groups of
        # rows.
        if (
                cnf_style.evenHBand in ["1", "true"] and
                self._tbl_look.noHBand in ["0", "false"]
        ):
            applied_styles.append(self._ref_cnd_styles.band2Horz)

        # get conditional style for intermediate rows, if the current row is
        # neither the first nor the last row of the table and the table has
        # an intermediate row style defined. This applies to odd groups of
        # rows.
        if (
                cnf_style.oddHBand in ["1", "true"] and
                self._tbl_look.noHBand in ["0", "false"]
        ):
            applied_styles.append(self._ref_cnd_styles.band1Horz)

        # get conditional style for last row, if the current row is the last
        # row of the table and the table has a last row style defined
        if (
                cnf_style.lastRow in ["1", "true"] and
                self._tbl_look.lastRow in ["1", "true"]
        ):
            applied_styles.append(self._ref_cnd_styles.lastRow)

        # get conditional style for first column, if the current column is the
        # first column of the table and the table has a first column style
        # defined
        if (
                cnf_style.firstColumn in ["1", "true"] and
                self._tbl_look.firstColumn in ["1", "true"]
        ):
            applied_styles.append(self._ref_cnd_styles.firstCol)

        # get conditional style for intermediate columns, if the current column
        # is neither the first nor the last column of the table and the table
        # has an intermediate column style defined. This applies to even
        # groups of columns.
        if (
                cnf_style.evenVBand in ["1", "true"] and
                self._tbl_look.noVBand in ["0", "false"]
        ):
            applied_styles.append(self._ref_cnd_styles.band2Vert)

        # get conditional style for intermediate columns, if the current column
        # is neither the first nor the last column of the table and the table
        # has an intermediate column style defined. This applies to odd
        # groups of columns.
        if (
                cnf_style.oddVBand in ["1", "true"] and
                self._tbl_look.noVBand in ["0", "false"]
        ):
            applied_styles.append(self._ref_cnd_styles.band1Vert)

        if (
                cnf_style.lastColumn in ["1", "true"] and
                self._tbl_look.lastColumn in ["1", "true"]
        ):
            applied_styles.append(self._ref_cnd_styles.lastCol)

        if len(applied_styles) == 0:
            return styles.TableCellProperty.init_zero()

        assert len(applied_styles) == 1

        tbl_prop = applied_styles[0].table_property
        tbl_cell_prop = applied_styles[0].table_cell_property

        # consolidate into table cell property; if a property is not defined
        # in the cell property, it is taken from the table property. Both
        # consist of shading and border properties.
        border_props = styles.TableCellBorders(
            **{
                b: getattr(tbl_cell_prop.tc_borders, b)
                for b in tbl_cell_prop.tc_borders.__slots__
            }
        )

        for b in tbl_prop.tbl_borders.__slots__:
            if getattr(border_props, b) is not None:
                continue
            setattr(border_props, b, getattr(tbl_prop.tbl_borders, b))

        shd_props = styles.ShdAttributes(
            **{
                s: getattr(tbl_cell_prop.tc_shd, s)
                for s in tbl_cell_prop.tc_shd.__slots__
            }
        )

        for s in tbl_prop.tbl_shd.__slots__:
            if getattr(shd_props, s) is not None:
                continue
            setattr(shd_props, s, getattr(tbl_prop.tbl_shd, s))

        return styles.TableCellProperty(
            tc_borders=border_props,
            tc_shd=shd_props
        )

    def __get_applied_tc_borders(
            self, tc: CT_Tc, cell_loc: str
    ) -> styles.TableCellBorders:
        r""" Finds the applied table cell borders for a given table cell by
        traversing through the hierarchy of styles defined for the
        table and its cells.

        @param tc: table cell element

        @return: table cell borders
        """
        # init table cell borders with None values
        tbl_cell_borders = styles.TableCellBorders.init_zero()
        tc_etree = etree.fromstring(tc.xml)

        # 1) parse cell level borders (source: word/document.xml)
        if (tc_pr := tc_etree.find(qn('w:tcPr'))) is not None:
            tc_pr_tc_borders = parse_tc_borders_element(
                tc_borders_elem=tc_pr.find(qn('w:tcBorders'))
            )

            tbl_cell_borders = fill_border_style(
                tbl_cell_borders, tc_pr_tc_borders
            )

        # 2) parse table level border styles (source: word/document.xml)
        # iterate through all table border elements (top, left, bottom,
        # right, insideH, insideV)
        tbl_cell_borders = fill_border_style(
            tbl_cell_borders,
            new_borders=convert_table_borders_to_cell_borders(
                cell_loc=cell_loc, tbl_borders=self._tbl_borders
            )
        )

        # 3) parse table style level properties (source: word/styles.xml) which
        # are always applied, though they can be overwritten by conditional
        # formatting, or by cell level properties
        tbl_cell_borders = fill_border_style(
            tbl_cell_borders,
            new_borders=convert_table_borders_to_cell_borders(
                cell_loc=cell_loc, tbl_borders=self._ref_tbl_props.tbl_borders
            )
        )

        # 4) parse table style cell border styles (source: word/styles.xml)
        # which are always applied, though they can be overwritten by
        # conditional formatting, or by cell level properties
        # from word/document.xml
        tbl_cell_borders = fill_border_style(
            tbl_cell_borders, self._ref_tc_props.tc_borders
        )

        return tbl_cell_borders

    def __get_applied_tc_shd(self, tc: CT_Tc) -> styles.ShdAttributes:
        r""" Finds the applied table cell shading for a given table cell by
        traversing through the hierarchy of styles defined for the
        table and its cells.

        @param tc: table cell element

        @return: table cell shading
        """

        # init table cell borders with None values
        tc_shd = styles.ShdAttributes.init_zero()
        tc_etree = etree.fromstring(tc.xml)

        # 1) parse cell level shading (source: word/document.xml)
        if (tc_pr := tc_etree.find(qn('w:tcPr'))) is not None:
            tc_pr_tc_shd = parse_shd_element(
                shd_elem=tc_pr.find(qn('w:shd'))
            )
            tc_shd = fill_shd_style(tc_shd, tc_pr_tc_shd)

        # 2) parse table level shading (source: word/document.xml)
        tc_shd = fill_shd_style(tc_shd, self._tbl_shading)

        # 3) parse table style level properties (source: word/styles.xml) which
        # are always applied, though they can be overwritten by conditional
        # formatting, or by cell level properties
        tc_shd = fill_shd_style(tc_shd, self._ref_tbl_props.tbl_shd)

        # 4) parse table style cell shading (source: word/styles.xml) which are
        # always applied, though they can be overwritten by conditional
        # formatting, or by cell level properties from word/document.xml
        tc_shd = fill_shd_style(tc_shd, self._ref_tc_props.tc_shd)

        return tc_shd

    def __parse_table_properties(self):
        r"""This function parses the table properties stored in
        word/document.xml and stores the following information in the
        corresponding class attributes:

        - table borders: table borders are stored in the class attribute
            _tbl_borders and are of type styles.TableBorders

        - table shading: table shading is stored in the class attribute
            _tbl_shading and is of type styles.ShdAttributes

        - table look: table look is stored in the class attribute _tbl_look
            and is of type styles.TableLook
        """
        tbl_pr_etree = etree.fromstring(self._ct_tbl.tblPr.xml)

        # table borders element
        self._tbl_borders = parse_tbl_borders_element(
            tbl_borders_elem=tbl_pr_etree.find(qn('w:tblBorders'))
        )

        # table shading element
        self._tbl_shading = parse_shd_element(
            shd_elem=tbl_pr_etree.find(qn('w:shd'))
        )

        # get tbl look element: this element determines in which regions the
        # style defined in the referenced style is applied.
        self._tbl_look = parse_tbl_look_element(
            tbl_look_elem=tbl_pr_etree.find(qn('w:tblLook'))
        )

    def __parse_tbl_ref_style(self):
        r""" This function parses the table style properties stored in the
        referenced style from word/styles.xml. It stores the following
        information in the corresponding class attributes:

        - conditional styles: conditional styles are stored in the class
            attribute _ref_cnd_styles and are of type styles.ConditionalStyles

        - table properties: table properties are stored in the class attribute
            _ref_tbl_props and are of type styles.TableProperty. These
            properties are always applied, though they can be overwritten by
            conditional formatting, or by cell level properties.

        - table cell properties: table cell properties are stored in the class
            attribute _ref_tc_props and are of type styles.TableCellProperty.
            These properties are always applied, though they can be overwritten
            by conditional formatting, or by cell level properties defined in
            word/document.xml.

        """
        try:
            style_etree = etree.fromstring(self._ct_tbl_ref_style.xml)
        except AttributeError:
            style_etree = None

        # parse conditional styles
        self._ref_cnd_styles = parse_ref_conditional_styles(style_etree)

        # parse table properties which are not conditional and are always
        # applied (though they may be overridden by conditional formatting)
        self._ref_tbl_props = styles.TableProperty.init_zero()
        if style_etree is not None:
            self._ref_tbl_props = parse_tbl_pr_element(
                tbl_pr=style_etree.find(qn('w:tblPr'))
            )

        # parse table cell properties which are not conditional and are always
        # applied (though they may be overridden by conditional formatting)
        self._ref_tc_props = styles.TableCellProperty.init_zero()
        if style_etree is not None:
            self._ref_tc_props = parse_tc_pr_element(
                tc_pr=style_etree.find(qn('w:tcPr'))
            )
