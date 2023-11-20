from docx.document import Document as _Document

from src.annotation.colorization import ColorizationHandler
from src.annotation.colorization.entities.tables import colorize_table

import settings


def colorize_header_and_footer(
        document: _Document, colorization_handler: ColorizationHandler
):
    r""" Colorize header and footer of the document.

    @param document: the document to colorize
    @param colorization_handler: the colorization handler to use for
        colorization of the header and footer
    """
    # colorize document header
    header_name = settings.colors.get_entity_name(settings.colors.COLOR_HEADER)
    _colorize(
        document, entity_name=header_name,
        colorization_handler=colorization_handler
    )

    # colorize document footer
    footer_name = settings.colors.get_entity_name(settings.colors.COLOR_FOOTER)
    _colorize(
        document, entity_name=footer_name,
        colorization_handler=colorization_handler
    )


def _colorize(
        document: _Document,
        entity_name: str,
        colorization_handler: ColorizationHandler
):
    r""" Colorize header or footer of the document.

    @param document: the document to colorize
    @param entity_name: either "header" or "footer"
    @param colorization_handler: the colorization handler to use for
        colorization of the header and footer
    """
    assert entity_name in [
        settings.entities.ENTITY_HEADER_NAME,
        settings.entities.ENTITY_FOOTER_NAME
    ]

    color = settings.colors.ENTITY_NAME_TO_COLOR[entity_name]

    for section in document.sections:
        header_or_footer_obj = getattr(section, entity_name)

        # skip if obj is linked to previous section
        if header_or_footer_obj.is_linked_to_previous:
            continue

        # colorize paragraphs
        for par in header_or_footer_obj.paragraphs:
            if len(par.text) == 0:
                continue

            colorization_handler.assign_par_color(
                par=par,
                base_color=color,
                decision_source=settings.annotation.ANNOTATION_BUILTIN
            )

        # colorize tables as footer / header
        for table in header_or_footer_obj.tables:
            colorize_table(
                table=table,
                base_color_table=color,
                base_color_table_header=color,
                colorization_handler=colorization_handler,
                sat_val_step=0
            )
